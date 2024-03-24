# This file aims to map the shen parcellation to the AAL atlas
# The mapping is based on the overlap of the two atlases

from atlasTransform.atlasTransform.utils.atlas import load_shen_268
import nibabel as nib
from nilearn import datasets, plotting
import numpy as np
import pandas as pd
from scipy.stats import mode
from scipy.spatial.distance import cdist
from plotting_preparation import (new_df, new_df_fc_movie, new_df_effect_of_movie, new_df_fc_effect_of_movie,
                                  new_df_rest_last_60_TR, new_df_fc_rest_last_60_TR,
                                  new_df_combined, new_df_fc_combined, new_df_fc_everything, new_df_movie_everything)
from brain_plotting import overlap_mask, movie_new, rest_new
from docx import Document

# Load the shen parcellation
shen_atlas = load_shen_268(2)
shen_data = shen_atlas.get_fdata()

# Load the AAL atlas
aal_atlas = datasets.fetch_atlas_aal(version='SPM12')
aal_img = nib.load(aal_atlas.maps)
aal_data = aal_img.get_fdata()

# map indices to labels
aal_labels = aal_atlas.labels
aal_indices = aal_atlas.indices
index_to_name = dict(zip(aal_indices, aal_labels))

# Get the unique labels in the shen parcellation
shen_labels = np.unique(shen_data)[1:]

# Identify unique labels (ignoring zero which often denotes background)
aal_coords_label = np.unique(aal_data)[1:]

# Initialize an empty list to store the centroid coordinates
aal_coords = []

# Loop through each unique label and calculate the centroid
for label in aal_coords_label:
    coords = np.argwhere(aal_data == label)
    centroid = np.mean(coords, axis=0)
    aal_coords.append(centroid)

# Convert the list to a NumPy array for numerical operations
aal_coords = np.array(aal_coords)

# Initialize a dictionary to hold mapping from Shen to AAL
shen_to_aal_mapping = {}

# Loop over each region in Shen parcellation
for shen_label in shen_labels:
    # Get AAL labels for all voxels in the current Shen parcel
    aal_labels_in_shen = aal_data[shen_data == shen_label]

    # Ensure non-zero AAL labels are considered (assuming 0 is the background label)
    aal_labels_in_shen = aal_labels_in_shen[aal_labels_in_shen != 0]

    # Get the most common AAL label in the current Shen parcel (Majority Voting)
    most_common_aal_label, count = mode(aal_labels_in_shen, keepdims=True)

    # Check if there's valid mapping (i.e., the Shen parcel has non-zero AAL labels)
    if count > 0:
        aal_label_name = index_to_name[str(int(most_common_aal_label))]
        shen_to_aal_mapping[shen_label] = aal_label_name
    else:
        shen_coords = np.argwhere(shen_data == shen_label)
        centroid = np.mean(shen_coords, axis=0)
        distances = cdist([centroid], aal_coords)
        nearest_index = np.argmin(distances)
        nearest_aal_label = aal_coords_label[nearest_index]
        aal_label_name = index_to_name[str(int(nearest_aal_label))]
        shen_to_aal_mapping[shen_label] = aal_label_name

df_mapping = pd.DataFrame(list(shen_to_aal_mapping.items()), columns=['Node Number', 'AAL Label'])


# now we have a dictionary mapping from Shen to AAL
# we can use this to map the Shen parcellation to AAL
def single_z_score_table(df):
    df_merged = pd.concat([df_mapping, df], axis=1)
    df_filtered = df_merged[~df_merged['u1'].isna()]
    df_filtered.loc[:, 'Node Number'] = df_filtered['Node Number'].astype(int).astype(str)
    df_filtered.loc[:, 'u1'] = np.round(df_filtered['u1'], 3)
    # reverse the sign of the z-score
    df_filtered.loc[:, 'u1'] = df_filtered['u1'] * -1
    df_filtered.loc[:, 'u1'] = df_filtered['u1'].apply(lambda x: "{:.3f}".format(x))
    df_filtered.loc[:, 'boot_ratio'] = np.round(df_filtered['boot_ratio'], 3)
    df_filtered.loc[:, 'boot_ratio'] = df_filtered['boot_ratio'] * -1
    df_filtered.loc[:, 'boot_ratio'] = df_filtered['boot_ratio'].apply(lambda x: "{:.3f}".format(x))
    df_final = df_filtered.rename(columns={'u1': 'PLS Loading', 'boot_ratio': 'Z-score'})

    return df_final


def aggregated_z_score_table(*dfs, table_name):

    tables = []
    for df in dfs:
        table = single_z_score_table(df)
        tables.append(table)

    # Resetting index before concatenation
    tables = [table.reset_index(drop=True) for table in tables]

    table_final = pd.concat(tables, axis=1)

    # export the mapping as a table
    # Create a new Word Document
    doc = Document()

    # Add a table to the Word document
    # The table will have df.shape[0] rows and df.shape[1] columns
    table = doc.add_table(rows=table_final.shape[0] + 1, cols=table_final.shape[1])

    # Insert the column names
    for i, column in enumerate(table_final.columns):
        table.cell(0, i).text = column

    # Insert the values of dataframe
    for i in range(table_final.shape[0]):
        for j in range(table_final.shape[1]):
            table.cell(i + 1, j).text = str(table_final.values[i, j])

    # Save the doc
    doc.save(table_name + '.docx')


# aggregated_z_score_table(new_df, new_df_fc_movie, 'decoupled_z_score_table')
# aggregated_z_score_table(new_df_effect_of_movie, new_df_fc_effect_of_movie, 'effect_of_movie_z_score_table')
# aggregated_z_score_table(new_df_rest_last_60_TR, new_df_fc_rest_last_60_TR, 'effect_of_propofol_z_score_table')
# aggregated_z_score_table(new_df_combined, new_df_fc_combined, 'combined_effect_z_score_table')
# aggregated_z_score_table(new_df_movie_everything, new_df_fc_everything, 'movie_everything_z_score_table')

# define a function to change the data format
def change_format(list):
    # change to dataframe
    df = pd.DataFrame(list, columns=['u1'])
    # add a column of boot_ratio
    df['boot_ratio'] = 0

    return df


# change the format of the data
df_hurst_effect_of_movie = change_format(movie_new)
df_hurst_effect_of_propofol = change_format(rest_new)

# create a table for the z-scores
aggregated_z_score_table(df_hurst_effect_of_movie, df_hurst_effect_of_propofol, table_name='effect_of_movie_z_score_table1')


def aggregated_z_score_table1(df_hurst, df_fc, table_name):
    hurst_table = single_z_score_table(df_hurst)
    fc_table = single_z_score_table(df_fc)

    # Resetting index before concatenation
    hurst_table = hurst_table.reset_index(drop=True)
    fc_table = fc_table.reset_index(drop=True)

    table_final = pd.concat([hurst_table, fc_table], axis=1)

    # export the mapping as a table
    # Create a new Word Document
    doc = Document()

    # Add a table to the Word document
    # The table will have df.shape[0] rows and df.shape[1] columns
    table = doc.add_table(rows=table_final.shape[0] + 1, cols=table_final.shape[1])

    # Insert the column names
    for i, column in enumerate(table_final.columns):
        table.cell(0, i).text = column

    # Insert the values of dataframe
    for i in range(table_final.shape[0]):
        for j in range(table_final.shape[1]):
            table.cell(i + 1, j).text = str(table_final.values[i, j])

    # Save the doc
    doc.save(table_name + '.docx')


aggregated_z_score_table1(df_hurst_effect_of_movie, df_hurst_effect_of_propofol,
                             table_name='effect_of_movie_z_score_table')
